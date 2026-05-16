"""Convert README.md → readme.docx with real hyperlinks and embedded screenshots.

Usage:
    python docs/scripts/make_docx.py
Output:
    synapse-docker/readme.docx
"""

import os, re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
README   = os.path.join(BASE_DIR, "README.md")
OUTPUT   = os.path.join(BASE_DIR, "readme.docx")

HYPERLINK_REL = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
)

doc = Document()

# ── global style tweaks ───────────────────────────────────────────────────────
for lvl in range(1, 5):
    doc.styles[f"Heading {lvl}"].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

# ── regex ─────────────────────────────────────────────────────────────────────
HTML_IMG_RE = re.compile(r'<img\b[^>]*\bsrc="([^"]+)"', re.I)
MD_IMG_RE   = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")

_TOKEN_RE = re.compile(
    r"(\*\*(?:[^*]|\*(?!\*))+\*\*"   # **bold** (may contain links)
    r"|`[^`]+`"                       # `code`
    r"|!\[[^\]]*\]\([^)]+\)"          # ![img](url)
    r"|\[[^\]]+\]\([^)]+\)"           # [label](url) hyperlink
    r")"
)
_LINK_RE    = re.compile(r"^\[([^\]]+)\]\(([^)]+)\)$")
_IMG_RE     = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)$")

# ── helpers ───────────────────────────────────────────────────────────────────

def resolve(src: str) -> str | None:
    path = os.path.join(BASE_DIR, src.replace("/", os.sep))
    return path if os.path.exists(path) else None


def add_image_to_doc(doc: Document, src: str, width=Inches(5.5)):
    path = resolve(src)
    if path:
        try:
            doc.add_picture(path, width=width)
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            doc.add_paragraph(f"[Image: {src}]").italic = True
    else:
        doc.add_paragraph(f"[Image not found: {src}]").italic = True


def add_image_grid(doc: Document, srcs: list, cols: int = 2):
    """Insert images into a borderless cols-wide table (2×2 for 4 images)."""
    rows_needed = (len(srcs) + cols - 1) // cols
    tbl = doc.add_table(rows=rows_needed, cols=cols)

    # Remove all table and cell borders
    tbl_xml = tbl._tbl
    tblPr = tbl_xml.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_xml.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    tblPr.append(tblBorders)

    cell_width = Inches(3.0)  # two columns fit within 6.5" content area

    for idx, src in enumerate(srcs):
        cell = tbl.cell(idx // cols, idx % cols)
        path = resolve(src)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if path:
            try:
                p.add_run().add_picture(path, width=cell_width)
            except Exception:
                p.add_run(f"[Image: {src}]").italic = True
        else:
            p.add_run(f"[Image not found: {src}]").italic = True

    doc.add_paragraph()


def add_hyperlink(para, url: str, text: str, bold: bool = False):
    """Insert a real clickable hyperlink run into para."""
    r_id = para.part.relate_to(url, HYPERLINK_REL, is_external=True)

    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Explicit blue + underline (reliable across all Word versions)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    if bold:
        rPr.append(OxmlElement("w:b"))

    run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    # preserve leading/trailing spaces
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    run.append(t)

    hl.append(run)
    para._p.append(hl)
    return hl


def apply_inline(para, text: str, bold: bool = False):
    """Parse bold / inline-code / [label](url) links; produce runs and hyperlinks.

    Bold tokens are handled recursively so **[link](url)** emits a bold hyperlink.
    """
    for part in _TOKEN_RE.split(text):
        if not part:
            continue

        # ── **bold** ── recurse so nested links inside bold work
        if part.startswith("**") and part.endswith("**"):
            apply_inline(para, part[2:-2], bold=True)

        # ── `inline code` ──
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.bold = bold
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)

        # ── ![img](url) — show alt text inline ──
        elif part.startswith("!["):
            m = _IMG_RE.match(part)
            if m and m.group(1):
                run = para.add_run(f"[{m.group(1)}]")
                run.italic = True

        # ── [label](url) — real hyperlink ──
        elif part.startswith("["):
            m = _LINK_RE.match(part)
            if m:
                add_hyperlink(para, m.group(2), m.group(1), bold=bold)

        # ── plain text ──
        else:
            run = para.add_run(part)
            run.bold = bold


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_table(doc: Document, md_lines: list):
    rows = [l for l in md_lines if not re.match(r"^\|[-| :]+\|$", l.strip())]
    if not rows:
        return
    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        parsed.append(cells)

    ncols = max(len(r) for r in parsed)
    table = doc.add_table(rows=0, cols=ncols)
    table.style = "Table Grid"

    for i, row in enumerate(parsed):
        while len(row) < ncols:
            row.append("")
        tr = table.add_row()
        for j, cell_text in enumerate(row):
            cell = tr.cells[j]
            p = cell.paragraphs[0]
            p.clear()
            apply_inline(p, cell_text, bold=(i == 0))
            if i == 0:
                set_cell_bg(cell, "D6E4F0")
    doc.add_paragraph()


def add_code_block(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F0F0F0")
    pPr.append(shd)


# ── main parser ────────────────────────────────────────────────────────────────

with open(README, encoding="utf-8") as f:
    lines = f.readlines()

i = 0
table_buf: list = []
in_code   = False
code_buf: list  = []

while i < len(lines):
    line = lines[i].rstrip("\n")

    # ── fenced code block ──
    if line.startswith("```"):
        if not in_code:
            in_code  = True
            code_buf = []
        else:
            add_code_block(doc, "\n".join(code_buf))
            in_code = False
        i += 1
        continue

    if in_code:
        code_buf.append(line)
        i += 1
        continue

    # ── <p ...> block: collect images from all consecutive <p> blocks ──
    if re.match(r"^\s*<p\b", line, re.I):
        all_srcs = []
        while i < len(lines):
            chunk = lines[i].rstrip("\n")
            for src in HTML_IMG_RE.findall(chunk):
                all_srcs.append(src)
            i += 1
            if re.match(r"^\s*</p>", chunk, re.I):
                # peek ahead past blank lines for another <p> block
                j = i
                while j < len(lines) and not lines[j].strip():
                    j += 1
                if j < len(lines) and re.match(r"^\s*<p\b", lines[j], re.I):
                    i = j  # jump to next <p> and keep collecting
                else:
                    break
        if all_srcs:
            add_image_grid(doc, all_srcs)
        continue

    # ── standalone <img> tag ──
    if re.match(r"^\s*<img\b", line, re.I):
        m = HTML_IMG_RE.search(line)
        if m:
            add_image_to_doc(doc, m.group(1))
        i += 1
        continue

    # ── <details> / <summary> / </details> ──
    if re.match(r"^\s*<(/?)(?:details|summary)\b", line, re.I):
        m = re.match(r"^\s*<summary>(.*?)</summary>", line, re.I)
        if m and m.group(1).strip():
            p = doc.add_paragraph()
            run = p.add_run(m.group(1).strip())
            run.bold = True
            run.italic = True
        i += 1
        continue

    # ── horizontal rule ──
    if re.match(r"^---+\s*$", line):
        doc.add_paragraph("─" * 60)
        i += 1
        continue

    # ── table rows ──
    if line.startswith("|"):
        table_buf.append(line)
        i += 1
        continue
    elif table_buf:
        add_table(doc, table_buf)
        table_buf = []
        # fall through to process current line

    # ── heading ──
    m = re.match(r"^(#{1,4})\s+(.*)", line)
    if m:
        level = len(m.group(1))
        heading_text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", m.group(2).strip())
        doc.add_heading(heading_text, level=level)
        i += 1
        continue

    # ── blockquote: collect all consecutive > lines into one paragraph ──
    if line.startswith(">"):
        parts = []
        while i < len(lines) and lines[i].startswith(">"):
            parts.append(lines[i].rstrip("\n").lstrip("> ").strip())
            i += 1
        content = " ".join(parts)
        style_names = [s.name for s in doc.styles]
        p = doc.add_paragraph(
            style="Intense Quote" if "Intense Quote" in style_names else "Normal"
        )
        apply_inline(p, content)
        continue

    # ── standalone markdown image ──
    img_m = MD_IMG_RE.match(line.strip())
    if img_m:
        add_image_to_doc(doc, img_m.group(2))
        i += 1
        continue

    # ── unordered list ──
    m = re.match(r"^(\s*)[-*]\s+(.*)", line)
    if m:
        indent = len(m.group(1)) // 2
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25 * (indent + 1))
        apply_inline(p, m.group(2))
        i += 1
        continue

    # ── ordered list ──
    m = re.match(r"^(\s*)\d+\.\s+(.*)", line)
    if m:
        p = doc.add_paragraph(style="List Number")
        apply_inline(p, m.group(2))
        i += 1
        continue

    # ── blank line ──
    if not line.strip():
        i += 1
        continue

    # ── normal paragraph ──
    p = doc.add_paragraph()
    apply_inline(p, line.strip())
    i += 1

if table_buf:
    add_table(doc, table_buf)

doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
